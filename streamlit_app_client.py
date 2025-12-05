# streamlit_app_client.py
# Website Accessibility Checker — Client Version
# Identical scanner + plain-language client reports (.docx) + new-tab downloads

import asyncio
import datetime
import io
import json
import os
import urllib.parse
from collections import Counter, defaultdict
import base64

import pandas as pd
import streamlit as st
from jinja2 import Template
from playwright.async_api import async_playwright
from docx import Document
from auth.auth_module import AuthManager, check_authentication

# Load favicon - MUST be before any other Streamlit commands
try:
    from PIL import Image
    # Try multiple possible paths for the icon
    icon_paths = [
        "Assets/RN_Web_A11y_IconDesign Wrapped.png",
        "assets/RN_Web_A11y_IconDesign Wrapped.png",
        os.path.join(os.path.dirname(__file__), "Assets", "RN_Web_A11y_IconDesign Wrapped.png"),
    ]
    
    favicon = None
    for path in icon_paths:
        try:
            favicon = Image.open(path)
            print(f"✅ Loaded favicon from: {path}")
            break
        except FileNotFoundError:
            continue
        except Exception as e:
            print(f"⚠️ Error loading {path}: {e}")
            continue
    
    if favicon:
        st.set_page_config(
            page_title="Website Accessibility Checker (Client Files)",
            page_icon=favicon,
            layout="wide"
        )
    else:
        st.set_page_config(
            page_title="Website Accessibility Checker (Client Files)",
            page_icon="♿",
            layout="wide"
        )
except ImportError:
    # Pillow not installed, use emoji
    st.set_page_config(
        page_title="Website Accessibility Checker (Client Files)",
        page_icon="♿",
        layout="wide"
    )
except Exception as e:
    print(f"⚠️ Error setting page config: {e}")
    st.set_page_config(
        page_title="Website Accessibility Checker (Client Files)",
        page_icon="♿",
        layout="wide"
    )

# -------------------- CONFIG --------------------
DEFAULT_MAX_PAGES = 40
SEVERITY_MAP = {
    "critical": "Critical",
    "serious": "Serious",
    "moderate": "Moderate",
    "minor": "Minor",
    None: "Moderate",
}
CATEGORY_BY_CRITERION = [
    ("Page Structure and Headings", ["heading", "title", "landmark", "structure"]),
    ("Keyboard Navigation", ["keyboard", "focus"]),
    ("Text and Colour Contrast", ["contrast"]),
    ("Images and Alt Text", ["image", "alt", "non-text"]),
    ("Forms and Labels", ["label", "form", "error"]),
    ("Mobile Accessibility", ["target", "viewport", "reflow"]),
    ("Screen Reader Experience", ["aria", "role", "link-name", "button-name"]),
]

STATEMENT_TEMPLATE = """
# Accessibility statement for {{ org_name }}

**Status**
This website aims to conform to **WCAG {{ wcag_version }} Level AA**.
Based on our latest audit ({{ scanned_at }}), the site is **partially compliant**
due to the non-compliances listed below.

{% if groups %}
## Non-accessible content
{% for g in groups %}
### {{ g.category }} — {{ g.criterion }}
{% for v in g["items"] %}
- **Page**: {{ v.page }}
  - **Issue**: {{ v.summary }}
  - **Example**: `{{ v.selector }}`
  - **Impact**: {{ v.impact }}
{% endfor %}
{% endfor %}
{% else %}
No known issues. This site is fully compliant with WCAG {{ wcag_version }} Level AA.
{% endif %}

_Last updated: {{ scanned_at }}._
""".strip()

# -------------------- HELPERS --------------------
def bucket_name(text: str) -> str:
    t = (text or "").lower()
    for name, keys in CATEGORY_BY_CRITERION:
        if any(k in t for k in keys):
            return name
    return "Other"

# -------------------- SCANNER --------------------
async def scan_site(start_url: str, max_pages: int = DEFAULT_MAX_PAGES):
    # Try to ensure Playwright browsers are installed
    try:
        from playwright_setup import ensure_playwright_browsers
        ensure_playwright_browsers()
    except Exception as e:
        print(f"⚠️ Could not auto-install browsers: {e}")
    host = urllib.parse.urlparse(start_url).netloc
    results = {
        "scanned_at": datetime.date.today().isoformat(),
        "site": start_url,
        "pages_scanned": 0,
        "wcag_version": "2.2",
        "violations": [],
    }

    visited, queue = set(), [start_url]

    try:
        async with async_playwright() as p:
            # Try to launch browser - handle installation issues gracefully
            try:
                browser = await p.chromium.launch()
            except Exception as browser_error:
                error_msg = str(browser_error).lower()
                # Check if it's a browser installation issue
                if any(keyword in error_msg for keyword in ["executable", "browser", "not found", "doesn't exist"]):
                    st.error("❌ **Playwright browsers are not installed**")
                    st.warning("""
                    **This app requires Playwright browsers to scan websites.**
                    
                    **Quick Fix:** 
                    - Use the "🔧 Install Playwright Browsers" section in the sidebar
                    - Click "📦 Install Browsers Now" button
                    - Wait 2-3 minutes for installation to complete
                    
                    **If installation fails:**
                    - This may be a Streamlit Cloud limitation
                    - Contact Streamlit support about browser installation
                    - Or run the app locally: `playwright install chromium`
                    """)
                    st.info("💡 **Tip:** Check the sidebar for the browser installation button!")
                    return results
                else:
                    # Other errors - show them
                    st.error(f"❌ Browser launch failed: {str(browser_error)}")
                    return results
            
            context = await browser.new_context()
            page = await context.new_page()

            while queue and len(visited) < max_pages:
                url = queue.pop(0)
                if url in visited:
                    continue
                visited.add(url)

                try:
                    await page.goto(url, wait_until="domcontentloaded", timeout=30000)
                except Exception:
                    continue

                # enqueue same-site links
                try:
                    links = await page.eval_on_selector_all("a[href]", "els => els.map(e => e.href)")
                    for href in links:
                        if urllib.parse.urlparse(href).netloc == host and href not in visited and href not in queue:
                            queue.append(href)
                except Exception:
                    pass

                # run axe-core
                try:
                    await page.add_script_tag(url="https://cdn.jsdelivr.net/npm/axe-core@4.9.1/axe.min.js")
                    axe = await page.evaluate("() => axe.run(document, { resultTypes: ['violations'] })")
                except Exception:
                    axe = {"violations": []}

                for v in axe.get("violations", []):
                    selector = v.get("nodes", [{}])[0].get("target", ["?"])[0]
                    results["violations"].append({
                        "page": url,
                        "criterion": f"{v.get('id')} — {v.get('help')}",
                        "axe_id": v.get("id"),
                        "severity": v.get("impact"),
                        "selector": selector,
                        "description": v.get("description"),
                    })

                results["pages_scanned"] = len(visited)

            await browser.close()
    except Exception as e:
        st.error(f"❌ Error during site scan: {str(e)}")
        import traceback
        st.exception(e)
        return results

    return results

# -------------------- DATA SHAPING --------------------
def to_dataframe(results):
    rows = []
    for v in results.get("violations", []):
        rows.append({
            "Page": v.get("page", ""),
            "Criterion": v.get("criterion", ""),
            "Axe ID": v.get("axe_id", ""),
            "Severity": SEVERITY_MAP.get(v.get("severity"), "Moderate"),
            "Selector / Element": v.get("selector", ""),
            "Description": v.get("description", ""),
        })
    # Ensure DataFrame has expected columns even if empty
    return pd.DataFrame(rows, columns=["Page", "Criterion", "Axe ID", "Severity", "Selector / Element", "Description"])

def group_issues_plain(results):
    """Cluster violations into client-friendly buckets."""
    buckets = defaultdict(list)
    for v in results["violations"]:
        cat = bucket_name(v["criterion"])
        buckets[cat].append(v)

    summary = []
    for cat, items in buckets.items():
        sev_counts = Counter(SEVERITY_MAP.get(i["severity"], "Moderate") for i in items)
        pages = sorted(set(i["page"] for i in items))
        summary.append({
            "category": cat,
            "count": len(items),
            "severity_counts": sev_counts,
            "pages": pages,
            "examples": items[:3],
            "action": plain_action_for(cat),
        })
    summary.sort(key=lambda x: x["count"], reverse=True)
    return summary

def plain_action_for(category):
    lib = {
        "Page Structure and Headings": "Use one H1 per page, fix heading order, and add landmarks so screen readers can navigate easily.",
        "Keyboard Navigation": "Ensure all buttons and links are keyboard reachable and show a clear focus outline.",
        "Text and Colour Contrast": "Increase text contrast to at least 4.5:1 (AA).",
        "Images and Alt Text": "Add short, meaningful alt text; mark purely decorative images with empty alt (alt=\"\").",
        "Forms and Labels": "Add visible labels linked to each field; make error messages clear and programmatically connected to their fields.",
        "Mobile Accessibility": "Add a viewport meta tag, increase touch targets ~44×44 px, and verify menus work on phones.",
        "Screen Reader Experience": "Give controls clear accessible names and correct ARIA roles/states so they’re announced properly.",
    }
    return lib.get(category, "Update this area to meet WCAG AA guidance.")

def why_it_matters(category):
    lib = {
        "Page Structure and Headings": "Headings and landmarks help screen reader users understand and navigate the page.",
        "Keyboard Navigation": "Some people rely on a keyboard; visible focus shows where they are.",
        "Text and Colour Contrast": "Low contrast makes text difficult for people with low vision or colour blindness.",
        "Images and Alt Text": "Alt text lets screen readers describe images to people who can’t see them.",
        "Forms and Labels": "Clear labels and announced errors help everyone complete forms successfully.",
        "Mobile Accessibility": "Many users browse on phones; content and controls must work well on small screens.",
        "Screen Reader Experience": "Correct names/roles/states let assistive tech announce controls accurately.",
    }
    return lib.get(category, "This area affects how easily users can navigate and understand content.")

# -------------------- DOCX BUILDERS --------------------
def build_client_audit_docx(results, org_name):
    S = group_issues_plain(results)
    counts = Counter(SEVERITY_MAP.get(v["severity"], "Moderate") for v in results["violations"])

    doc = Document()
    doc.core_properties.title = "Accessibility Audit Report"

    doc.add_heading("Accessibility Audit Report", 0)
    doc.add_paragraph(f"Organisation: {org_name or urllib.parse.urlparse(results['site']).netloc}")
    doc.add_paragraph(f"Website: {results['site']}")
    doc.add_paragraph(f"Date: {results['scanned_at']}")

    doc.add_heading("1. Executive Summary", 1)
    doc.add_paragraph(
        f"We reviewed key pages against WCAG 2.2 AA. "
        f"{len(results['violations'])} issues were found across {results.get('pages_scanned', 0)} pages."
    )
    doc.add_paragraph("Severity breakdown:")
    for s in ["Critical", "Serious", "Moderate", "Minor"]:
        doc.add_paragraph(f"- {s}: {counts.get(s, 0)}")

    doc.add_heading("2. Key Findings", 1)
    for grp in S:
        doc.add_heading(f"{grp['category']} — {grp['count']} issues", 2)
        doc.add_paragraph(f"Why this matters: {why_it_matters(grp['category'])}")
        doc.add_paragraph(f"What to do: {grp['action']}")
        doc.add_paragraph("Examples:")
        for ex in grp["examples"]:
            doc.add_paragraph(f"• {ex['criterion']} — Page: {ex['page']} — Element: {ex['selector']}")

    doc.add_heading("3. Next Steps", 1)
    doc.add_paragraph("Prioritise Critical/Serious items first, then improve contrast, forms/labels, structure, and mobile experience.")

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf

def build_client_statement_docx(results, org_name, contact_name, contact_email):
    grouped = group_issues_plain(results)

    doc = Document()
    doc.core_properties.title = "Accessibility Statement"

    doc.add_heading("Accessibility Statement", 0)
    doc.add_paragraph(f"Website: {results['site']}")
    doc.add_paragraph(f"Organisation: {org_name or urllib.parse.urlparse(results['site']).netloc}")
    doc.add_paragraph(f"Last Updated: {results['scanned_at']}")

    doc.add_heading("Our Commitment", 1)
    doc.add_paragraph("We’re committed to making our website accessible to everyone and aim to meet WCAG 2.2 Level AA.")

    doc.add_heading("Current Status", 1)
    doc.add_paragraph("This website is partially compliant with WCAG 2.2 Level AA due to the areas for improvement listed below.")

    doc.add_heading("Areas for Improvement", 1)
    for grp in grouped:
        doc.add_heading(grp["category"], 2)
        doc.add_paragraph(f"{grp['count']} issues found. {plain_action_for(grp['category'])}")

    doc.add_heading("Planned Improvements", 1)
    doc.add_paragraph("We will address high-impact issues first and continue improvements as part of an ongoing accessibility roadmap.")

    doc.add_heading("Feedback and Contact", 1)
    doc.add_paragraph(f"If you need an alternative format or want to report a barrier, contact {contact_name} at {contact_email}.")

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf

# -------------------- STREAMLIT UI --------------------
st.set_page_config(
    page_title="Website Accessibility Checker (Client Files)",
    page_icon="♿",
    layout="wide"
)

# Initialize authentication
if 'auth_manager' not in st.session_state:
    st.session_state.auth_manager = AuthManager()

auth_manager = st.session_state.auth_manager

# Check authentication - show login if not authenticated
if not check_authentication(auth_manager):
    st.stop()

# User is authenticated - show main app
st.title("Website Accessibility Checker")
st.caption("Identical scanner + client-friendly report/statement downloads.")

with st.sidebar:
    # Logout button at top
    if st.button("🚪 Logout", use_container_width=True, key="client_logout_btn"):
        st.session_state.authenticated_user = None
        st.rerun()
    
    st.divider()
    
    # Browser installation section
    with st.expander("🔧 Install Playwright Browsers", expanded=False):
        st.caption("Required for website scanning. Click to install if you see browser errors.")
        if st.button("📦 Install Browsers Now", use_container_width=True, type="secondary", key="client_install_browsers"):
            with st.spinner("Installing Playwright browsers... This may take 2-3 minutes."):
                try:
                    from playwright_setup import install_playwright_browsers
                    success = install_playwright_browsers()
                    if success:
                        st.success("✅ Browsers installed successfully! You can now run scans.")
                        st.balloons()
                    else:
                        st.error("❌ Installation failed. Check the logs for details.")
                except Exception as e:
                    st.error(f"❌ Error: {str(e)}")
                    st.info("💡 If this fails, browsers may need to be installed via Streamlit Cloud settings.")
    
    st.divider()
    
    url = st.text_input(
        "Start URL (include https://)",
        placeholder="https://example.com",
        key="client_url_input"
    )
    max_pages = st.number_input(
        "Max pages to scan",
        min_value=1, max_value=500,
        value=DEFAULT_MAX_PAGES, step=1,
        key="client_max_pages"
    )
    org_name = st.text_input("Organisation name", "", key="client_org_name")
    contact_name = st.text_input("Contact name", "Accessibility Lead", key="client_contact_name")
    contact_email = st.text_input("Contact email", "accessibility@example.com", key="client_contact_email")
    run_btn = st.button("Run accessibility scan", type="primary", key="client_run_btn")

if run_btn:
    if not url or not url.startswith("http"):
        st.error("Please enter a valid URL starting with http or https.")
        st.stop()

    with st.spinner("Scanning… please wait (30–60 s)."):
        results = asyncio.run(scan_site(url, int(max_pages)))

    # Table for tech team
    df = to_dataframe(results)
    st.subheader("Summary")
    st.write(f"**Pages scanned:** {results.get('pages_scanned', 0)}")
    st.write(f"**Total issues found:** {len(results['violations'])}")
    if not df.empty:
        st.dataframe(df, use_container_width=True)
        st.download_button(
            "Download audit CSV",
            df.to_csv(index=False).encode("utf-8"),
            "a11y_audit.csv",
            "text/csv",
            key="csv_download_btn"
        )
    else:
        st.success("No issues found.")

    # Statement preview (on-screen)
    st.divider()
    st.header("Accessibility Statement (Preview)")
    preview_groups = []
    for g in group_issues_plain(results):
        preview_groups.append({
            "category": g["category"],
            "criterion": g["category"],
            "items": [
                {"page": ex["page"], "selector": ex["selector"],
                 "summary": ex["description"],
                 "impact": SEVERITY_MAP.get(ex["severity"], "Moderate")}
                for ex in g["examples"]
            ]
        })
    md = Template(STATEMENT_TEMPLATE).render(
        org_name=org_name or urllib.parse.urlparse(results["site"]).netloc,
        wcag_version=results["wcag_version"],
        scanned_at=results["scanned_at"],
        pages_scanned=results["pages_scanned"],
        groups=preview_groups,
    )
    st.markdown(md)

    # Client-friendly downloads (NEW-TAB links, no app reload)
    st.divider()
    st.header("Client-friendly downloads")

    audit_doc = build_client_audit_docx(results, org_name)
    stmt_doc  = build_client_statement_docx(results, org_name, contact_name, contact_email)

    def make_download_link(data_bytes, filename, label):
        b64 = base64.b64encode(data_bytes).decode()
        href = (
            f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;'
            f'base64,{b64}" download="{filename}" target="_blank">{label}</a>'
        )
        st.markdown(href, unsafe_allow_html=True)

    st.markdown("### 📄 Download your reports")
    make_download_link(audit_doc.getvalue(),
                       "Client_Accessibility_Audit_Report.docx",
                       "⬇️ Download Client Accessibility Audit Report (.docx)")
    make_download_link(stmt_doc.getvalue(),
                       "Client_Accessibility_Statement.docx",
                       "⬇️ Download Client Accessibility Statement (.docx)")

else:
    st.info("Enter a URL in the sidebar and click **Run accessibility scan**.")

